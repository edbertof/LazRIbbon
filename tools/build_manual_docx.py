from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
CONTENT_WIDTH_IN = 6.5


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.style = "Table Text"
    run = paragraph.add_run(text.strip())
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_table_width(table, widths_in: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = Inches(widths_in[min(idx, len(widths_in) - 1)])
            cell.width = width
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(round(widths_in[min(idx, len(widths_in) - 1)] * 1440)))


def clean_inline(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("__", "")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def add_markdown_runs(paragraph, text: str) -> None:
    token_re = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("["):
            label = re.match(r"\[([^\]]+)\]", token).group(1)
            run = paragraph.add_run(label)
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(8)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", 1)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(3)

    if "Caption Text" not in styles:
        caption = styles.add_style("Caption Text", 1)
    else:
        caption = styles["Caption Text"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(85, 85, 85)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)

    if "Table Text" not in styles:
        table_text = styles.add_style("Table Text", 1)
    else:
        table_text = styles["Table Text"]
    table_text.font.name = "Calibri"
    table_text.font.size = Pt(9)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.15


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(title)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(14)
    run = meta.add_run(subtitle)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(85, 85, 85)

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E8EEF5")
    set_cell_text(cell, "Target: LazRibbon 2.1.1 | Lazarus 4.8 | Generated from repository documentation")
    doc.add_paragraph()


def add_image(doc: Document, md_path: Path, alt: str, image_ref: str) -> None:
    image_path = (md_path.parent / image_ref).resolve()
    if not image_path.exists():
        doc.add_paragraph(f"[Image not found: {image_ref}]", style="Code Block")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))
    caption = doc.add_paragraph(alt, style="Caption Text")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    if col_count == 2:
        widths = [1.85, 4.65]
    elif col_count == 3:
        widths = [1.65, 2.2, 2.65]
    else:
        each = CONTENT_WIDTH_IN / col_count
        widths = [each] * col_count
    set_table_width(table, widths)
    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        for c_idx in range(col_count):
            text = clean_inline(row_data[c_idx]) if c_idx < len(row_data) else ""
            set_cell_text(row.cells[c_idx], text, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(row.cells[c_idx], "E8EEF5")
    doc.add_paragraph()


def parse_md_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
        raw = lines[i].strip().strip("|")
        cells = [cell.strip() for cell in raw.split("|")]
        if not all(set(cell.replace(":", "").strip()) <= {"-"} for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def build_docx(md_file: Path, out_file: Path, title: str, subtitle: str) -> None:
    doc = Document()
    configure_styles(doc)
    add_cover(doc, title, subtitle)

    lines = md_file.read_text(encoding="utf-8").splitlines()
    paragraph_buffer: list[str] = []
    in_code = False
    code_lang = ""

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer if part.strip())
        paragraph_buffer = []
        if not text:
            return
        p = doc.add_paragraph()
        add_markdown_runs(p, text)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            in_code = not in_code
            code_lang = stripped.strip("`").strip()
            if code_lang and in_code:
                p = doc.add_paragraph(style="Code Block")
                p.add_run(f"[{code_lang}]").bold = True
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph(style="Code Block")
            p.add_run(line)
            i += 1
            continue

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            flush_paragraph()
            add_image(doc, md_file, image_match.group(1), image_match.group(2))
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph()
            rows, new_i = parse_md_table(lines, i)
            add_table(doc, rows)
            i = new_i
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading_match:
            flush_paragraph()
            level = min(len(heading_match.group(1)), 3)
            text = clean_inline(heading_match.group(2))
            style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
            doc.add_paragraph(text, style=style)
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_markdown_runs(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            add_markdown_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("LazRibbon 2.1.1")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)

    out_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_file)


def main() -> int:
    targets = [
        (
            ROOT / "docs/manual/LAZRIBBON_MANUAL.md",
            ROOT / "docs/manual/LAZRIBBON_MANUAL.docx",
            "LazRibbon Manual",
            "Installation and illustrated usage guide",
        ),
        (
            ROOT / "docs/manual/LAZRIBBON_COMPONENT_REFERENCE.md",
            ROOT / "docs/manual/LAZRIBBON_COMPONENT_REFERENCE.docx",
            "LazRibbon Component Reference",
            "Property-by-property and event-by-event reference",
        ),
    ]
    for md_file, out_file, title, subtitle in targets:
        build_docx(md_file, out_file, title, subtitle)
        print(f"Wrote {out_file}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
